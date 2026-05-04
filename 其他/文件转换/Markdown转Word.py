# -*- coding: utf-8 -*-
"""
Markdown 转 Word 工具（支持 macOS / Windows / Linux）

将 .md 文件转换为 .docx 格式，适用于需求提炼、测试关键点分析等文档。

依赖：
  - 必选：python-docx  →  pip install python-docx
  - 可选：pandoc（转换质量更好）→  macOS: brew install pandoc

用法：
  单文件：  python3 Markdown转Word.py <md文件路径> [输出.docx]
  批量：    python3 Markdown转Word.py <文件夹路径>
  （Windows 下可将 python3 改为 python）

示例：
  python3 Markdown转Word.py 需求文档/需求理解输出/需求提炼分析.md
  python3 Markdown转Word.py 需求提炼分析.md 输出.docx
  python3 Markdown转Word.py 需求文档/需求理解输出/
"""
import atexit
import os
import re
import subprocess
import sys
from pathlib import Path

# 脚本结束时自动清理 __pycache__
_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _root not in sys.path:
    sys.path.insert(0, _root)
if os.path.join(_root, "其他", "系统维护") not in sys.path:
    sys.path.insert(0, os.path.join(_root, "其他", "系统维护"))
try:
    from 清理缓存 import clean_pycache
    atexit.register(lambda: clean_pycache(verbose=False))
except ImportError:
    pass

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


def _pandoc_available() -> bool:
    """检测系统是否已安装 pandoc（macOS 可 brew install pandoc）。"""
    try:
        subprocess.run(
            ["pandoc", "--version"],
            capture_output=True,
            check=True,
            timeout=5,
        )
        return True
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return False


def _convert_with_pandoc(md_path: Path, docx_path: Path) -> bool:
    """使用 pandoc 转换，成功返回 True，失败返回 False。"""
    try:
        subprocess.run(
            ["pandoc", str(md_path), "-o", str(docx_path), "--from", "markdown", "--to", "docx"],
            check=True,
            capture_output=True,
            timeout=60,
            cwd=str(md_path.parent),
        )
        return True
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return False


def parse_table_line(line: str) -> list:
    """解析表格行，返回单元格列表"""
    cells = []
    cell = ""
    in_pipe = False
    for i, c in enumerate(line):
        if c == '|':
            if in_pipe or i > 0:
                cells.append(cell.strip())
                cell = ""
            in_pipe = True
        else:
            cell += c
            in_pipe = False
    if cell.strip():
        cells.append(cell.strip())
    return [c for c in cells if c]  # 过滤空字符串


def is_table_separator(line: str) -> bool:
    """判断是否为表格分隔行 |---|---|"""
    if '|' not in line:
        return False
    parts = line.split('|')
    parts = [p.strip() for p in parts if p.strip()]
    return all(re.match(r'^[-:]+$', p) for p in parts)


def md_to_docx(md_path: str, docx_path: str = None, use_pandoc: bool = True) -> str:
    """
    将 Markdown 文件转换为 Word 文档。
    若 use_pandoc 为 True 且系统已安装 pandoc，优先使用 pandoc（质量更好）；否则用内置解析。

    :param md_path: Markdown 文件路径
    :param docx_path: 输出 Word 文件路径，默认与 md 同目录同主名
    :param use_pandoc: 是否优先使用 pandoc（macOS 建议 brew install pandoc）
    :return: 转换后的 docx 文件路径
    """
    md_file = Path(md_path).resolve()
    if not md_file.exists():
        raise FileNotFoundError(f"文件不存在: {md_path}")

    if docx_path is None:
        docx_path = md_file.with_suffix(".docx")
    else:
        docx_path = Path(docx_path).resolve()

    if use_pandoc and _pandoc_available():
        if _convert_with_pandoc(md_file, docx_path):
            return str(docx_path)

    return _md_to_docx_fallback(md_file, docx_path)


def _md_to_docx_fallback(md_file: Path, docx_path: Path) -> str:
    """使用 python-docx 解析 MD 并生成 docx（不依赖 pandoc）。"""
    doc = Document()

    with open(md_file, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    table_rows = []
    in_table = False
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        line_stripped = line.rstrip()

        # 代码块开始/结束
        if line_stripped.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_lines.append(line_stripped)
            i += 1
            continue

        # 空行
        if not line_stripped:
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # 表格分隔行
        if is_table_separator(line_stripped):
            i += 1
            continue

        # 表格行
        if line_stripped.startswith("|") and line_stripped.endswith("|"):
            cells = parse_table_line(line_stripped)
            if cells:
                table_rows.append(cells)
                in_table = True
            i += 1
            continue

        # 非表格行，先输出已缓存的表格
        if in_table and table_rows:
            _add_table_to_doc(doc, table_rows)
            table_rows = []
            in_table = False

        # 标题（# ~ #####）
        if line_stripped.startswith("##### "):
            doc.add_heading(line_stripped[6:].strip(), level=4)
        elif line_stripped.startswith("#### "):
            doc.add_heading(line_stripped[5:].strip(), level=3)
        elif line_stripped.startswith("### "):
            doc.add_heading(line_stripped[4:].strip(), level=2)
        elif line_stripped.startswith("## "):
            doc.add_heading(line_stripped[3:].strip(), level=1)
        elif line_stripped.startswith("# "):
            doc.add_heading(line_stripped[2:].strip(), level=0)
        elif line_stripped.strip() == "---":
            pass
        elif re.match(r"^[-*] ", line_stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _add_paragraph_with_format(p, line_stripped[2:])
        elif re.match(r"^\d+\. ", line_stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _add_paragraph_with_format(p, re.sub(r"^\d+\.\s*", "", line_stripped))
        else:
            _add_paragraph_with_format(doc.add_paragraph(), line_stripped)

        i += 1

    if table_rows:
        _add_table_to_doc(doc, table_rows)
    if in_code_block and code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    doc.save(str(docx_path))
    return str(docx_path)


def _add_table_to_doc(doc, rows: list):
    """将表格添加到文档"""
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    for ri, row_cells in enumerate(rows):
        for ci, cell_text in enumerate(row_cells):
            if ci < col_count:
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text
                # 表头行加粗
                if ri == 0 and len(rows) > 1:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True


def _add_paragraph_with_format(paragraph, text: str):
    """添加带格式的段落（支持 **粗体**、`行内代码`）。"""
    # 按 **bold** 和 `code` 切分并保持顺序
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif part:
            paragraph.add_run(part)


def convert_folder(folder: str, use_pandoc: bool = True) -> list:
    """批量转换文件夹内所有 .md 文件，返回生成的 docx 路径列表。"""
    folder_path = Path(folder).resolve()
    if not folder_path.is_dir():
        raise NotADirectoryError(f"不是文件夹: {folder}")
    results = []
    for f in sorted(folder_path.iterdir()):
        if f.suffix.lower() != ".md" or not f.is_file():
            continue
        try:
            out = md_to_docx(str(f), use_pandoc=use_pandoc)
            results.append(out)
            print(f"已生成: {out}")
        except Exception as e:
            print(f"跳过 {f.name}: {e}", file=sys.stderr)
    return results


def main():
    argv = [a for a in sys.argv[1:] if a != "--no-pandoc"]
    use_pandoc = "--no-pandoc" not in sys.argv

    if len(argv) < 1:
        print("用法: python3 Markdown转Word.py <md文件或文件夹> [输出.docx]")
        print("      python3 Markdown转Word.py --no-pandoc <md文件>  # 仅用内置解析")
        print("示例:")
        print("  python3 Markdown转Word.py 需求文档/需求理解输出/需求提炼分析.md")
        print("  python3 Markdown转Word.py 需求提炼分析.md 输出.docx")
        print("  python3 Markdown转Word.py 需求文档/需求理解输出/   # 批量")
        print("macOS 推荐安装 pandoc 以获得更好效果: brew install pandoc")
        sys.exit(1)

    target = Path(argv[0]).resolve()
    if target.is_dir():
        try:
            convert_folder(str(target), use_pandoc=use_pandoc)
        except Exception as e:
            print(f"批量转换失败: {e}", file=sys.stderr)
            sys.exit(1)
        return

    md_path = argv[0]
    docx_path = argv[1] if len(argv) >= 2 else None
    try:
        result = md_to_docx(md_path, docx_path, use_pandoc=use_pandoc)
        print(f"已生成: {result}")
    except Exception as e:
        print(f"转换失败: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
