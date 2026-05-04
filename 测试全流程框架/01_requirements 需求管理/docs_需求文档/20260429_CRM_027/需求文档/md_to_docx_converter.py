#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX 转换器
用于保持PRD文档的Markdown和DOCX格式版本同步
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

def parse_markdown(md_content):
    """解析Markdown内容，提取标题、段落、表格等元素"""
    elements = []
    lines = md_content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 处理标题
        if line.startswith('#'):
            level = line.count('#')
            title = line.lstrip('#').strip()
            elements.append({'type': 'heading', 'level': level, 'content': title})
            i += 1
        
        # 处理表格
        elif '|' in line and i + 1 < len(lines) and lines[i+1].count('-') >= 2:
            table_rows = []
            while i < len(lines) and '|' in lines[i]:
                cells = [c.strip() for c in lines[i].split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
            elements.append({'type': 'table', 'rows': table_rows})
        
        # 处理代码块
        elif line.startswith('```'):
            code_lang = line[3:].strip()
            code_content = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_content.append(lines[i])
                i += 1
            i += 1  # 跳过结束的```
            elements.append({'type': 'code', 'language': code_lang, 'content': '\n'.join(code_content)})
        
        # 处理列表项
        elif line.startswith('- ') or line.startswith('* '):
            list_items = []
            while i < len(lines) and (lines[i].startswith('- ') or lines[i].startswith('* ')):
                list_items.append(lines[i][2:].strip())
                i += 1
            elements.append({'type': 'list', 'items': list_items})
        
        # 处理普通段落
        elif line.strip():
            elements.append({'type': 'paragraph', 'content': line.strip()})
            i += 1
        
        else:
            i += 1
    
    return elements

def create_docx(elements, output_path):
    """根据解析的元素创建DOCX文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    for element in elements:
        if element['type'] == 'heading':
            level = element['level']
            content = element['content']
            
            if level == 1:
                heading = doc.add_heading(level=1)
                run = heading.add_run(content)
                run.font.size = Pt(16)
                run.font.bold = True
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                heading = doc.add_heading(level=2)
                run = heading.add_run(content)
                run.font.size = Pt(14)
                run.font.bold = True
            elif level == 3:
                heading = doc.add_heading(level=3)
                run = heading.add_run(content)
                run.font.size = Pt(12)
                run.font.bold = True
            else:
                heading = doc.add_heading(level=level)
                run = heading.add_run(content)
                run.font.size = Pt(11)
                run.font.bold = True
        
        elif element['type'] == 'paragraph':
            p = doc.add_paragraph(element['content'])
            p.paragraph_format.first_line_indent = Inches(0.25)
            p.paragraph_format.line_spacing = 1.5
        
        elif element['type'] == 'table':
            rows = element['rows']
            if not rows:
                continue
            
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for row_idx, row in enumerate(rows):
                for col_idx, cell in enumerate(row):
                    table.cell(row_idx, col_idx).text = cell
                    table.cell(row_idx, col_idx).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # 表头加粗
                    if row_idx == 0:
                        for paragraph in table.cell(row_idx, col_idx).paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
        elif element['type'] == 'code':
            p = doc.add_paragraph()
            run = p.add_run(element['content'])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
        
        elif element['type'] == 'list':
            for item in element['items']:
                p = doc.add_paragraph(item, style='List Bullet')
                p.paragraph_format.line_spacing = 1.5
    
    doc.save(output_path)
    print(f"✓ DOCX文件已生成: {output_path}")

def sync_versions(md_path, docx_dir):
    """同步Markdown和DOCX版本"""
    if not os.path.exists(md_path):
        print(f"✗ 错误：Markdown文件不存在: {md_path}")
        return
    
    # 提取版本号
    md_filename = os.path.basename(md_path)
    version = md_filename.split('_v')[-1].replace('.md', '')
    
    # 生成DOCX文件名
    docx_filename = md_filename.replace('.md', '.docx')
    docx_path = os.path.join(docx_dir, docx_filename)
    
    # 读取Markdown内容
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 解析并生成DOCX
    elements = parse_markdown(md_content)
    create_docx(elements, docx_path)
    
    return docx_path

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("用法: python md_to_docx_converter.py <markdown文件路径> <docx输出目录>")
        sys.exit(1)
    
    md_path = sys.argv[1]
    docx_dir = sys.argv[2]
    
    sync_versions(md_path, docx_dir)
